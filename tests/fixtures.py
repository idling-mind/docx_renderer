import pytest
import json
from pathlib import Path
from tempfile import TemporaryDirectory
from docx import Document

@pytest.fixture
def input_json():
    data = {
        "myvar1": "hello 1",
        "myvar2": "hello 2",
    }
    input_path = Path("input.json")
    input_path.write_text(json.dumps(data))
    yield input_path
    input_path.unlink()

@pytest.fixture
def input_template():
    """Input template for testing"""
    with TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir) / "template.docx"
        # Create a Word document with placeholders {{{myvar1}}} and {{{myvar2}}}
        doc = Document()
        doc.add_paragraph("{{{myvar1}}}")
        doc.add_paragraph("{{{myvar2}}}")
        doc.save(temp_path)
        yield temp_path