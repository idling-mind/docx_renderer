from __future__ import annotations
from pathlib import Path
from PIL import Image
from docx import Document
from docx.shared import Pt
from .exceptions import RenderError

def image(
    context: dict,
    width: int | None = None,
    height: int | None = None,
    remove_placeholder: bool = True,
):
    """Insert an image into the document at the location of the placeholder.

    Args:
        context (dict): Dictionary containing the following keys:
            - result: The result of evaluating the python statement.
            - paragraph: The docx paragraph object where the placeholder is present.
            - document: The output docx document object.
        preserve_aspect_ratio (bool, optional): Preserve the aspect ratio of the image.
            Defaults to True.
        remove_placeholder (bool, optional): Remove the placeholder after the image is inserted.
            Defaults to True.
        alignment (str, optional): Alignment of the image. Can be 'left', 'center', 'right'. Defaults to 'left'.
        crop (dict, optional): Dictionary containing the crop values for the image.
            The keys can be 'left', 'right', 'top', 'bottom'. Defaults to None.
            The values are given in percentage of the image size.
    """
    result = str(context["result"])
    paragraph = context["paragraph"]
    document = context["document"]

    if not Path(result).exists():
        raise RenderError(f"Image '{result}' not found.")

    if remove_placeholder:
        paragraph.text = ""

    paragraph.add_run().add_picture(result, width=width, height=height)

    # Optionally remove the placeholder text
def table(
    context: dict,
    first_row=True,
    remove_placeholder=True,
):
    """Insert a table into the document at the location of the placeholder.

    Args:
        context (dict): Dictionary containing the following keys:
            - result: The result of evaluating the python statement.
            - paragraph: The docx paragraph object where the placeholder is present.
            - document: The output docx document object.
        first_row (bool, optional): Show the first row as header. Defaults to True.
        remove_placeholder (bool, optional): Remove the placeholder after the table is inserted.
            Defaults to True.
    """
    result = context["result"]
    paragraph = context["paragraph"]
    document = context["document"]

    table_data = list(result)
    table = document.add_table(rows=len(table_data), cols=len(table_data[0]))

    for row_idx, row_data in enumerate(table_data):
        row = table.rows[row_idx]
        for col_idx, cell_data in enumerate(row_data):
            row.cells[col_idx].text = str(cell_data)
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)

    # Optionally remove the placeholder text
    if remove_placeholder:
        paragraph.text = ""
